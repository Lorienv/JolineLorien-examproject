##############################
#Text analyses of 1001-nights#
##############################


from os import listdir
# Make a corpus of the ten volumes:
def list_10_volumes(directory):
	volumes = []
	for volume in listdir(directory):
		if volume.startswith('arabian'):
			volumes.append(directory + '/' + volume)
	return volumes

corpus = list_10_volumes('data')
#print(corpus)

# How many characters does each volume in the corpus have? 
# and how many characters does the entire corpus have?
def calculate_characters(corpus):
	counter = 0
	characters_per_volume = []
	for volume in corpus:
		f = open(volume, 'rt', encoding='utf-8') 
		text = f.read()
		f.close()
		characters_per_volume.append(len(text))	
		counter+= len(text)
	return characters_per_volume, 'Together, the ten volumes have ' + str(counter) + ' characters.'

#print(calculate_characters(corpus))	

# How many lines does the entire corpus have?
def calculate_lines(corpus):
	count = 0
	for volume in corpus:
		f = open(volume, 'rt', encoding='utf-8') 
		for line in f:
			count += 1
		f.close()	
	#print('The corpus of ten volumes has ' + str(count) + ' lines.')	
			
#print(calculate_lines(corpus))	

# How many lines does each volume have?
def calculate_lines_II(file):
	count = 0
	f = open (file, 'rt', encoding='utf-8') 
	for line in f:
		count += 1
	f.close()
	return [file, ' has ', count, 'lines.']

for volume in corpus: 
	print(calculate_lines_II(volume))	

# In order to calculate the amount of words and sentences in each volume,
# I made a new corpus of the volumes using the nltk PlaintextCorpusReader
# which has some easy tools that can split a text into a list of words or sentences.

from nltk.corpus import PlaintextCorpusReader
corpus_root= 'data'
volumes = PlaintextCorpusReader(corpus_root, 'arabian.*')

list_of_sentences = volumes.sents()
#print('The ten volumes consist of ' + str(len(list_of_sentences)) + ' sentences')

list_of_words = volumes.words()
#print('The ten volumes consist of ' + str(len(list_of_words)) + ' words')

for item in volumes.fileids(): #calculate the amount of words in each volume
	print(item,':', len(volumes.words(item)), 'words')

for item in volumes.fileids(): #calculate the amount of sentences in each volume
	print(item,':', len(volumes.sents(item)), 'sentences')	

##################################################################
#visualisation of the statistiscs with basic plotting techniques
##################################################################

import matplotlib.pyplot as plt #pyhton library for plotting data
import numpy as np

#visualise the characters per volume
#make a list of the total characters per volume -
characters_per_volume = (calculate_characters(corpus))[0]
x = [1,2,3,4,5,6,7,8,9,10]
y = characters_per_volume
x_labels = ["vol1","vol2","vol3","vol4","vol5","vol6","vol7","vol8","vol9","vol10"]
y_labels = [719626, 713975, 740002, 608981, 805537, 620786, 816219, 728202, 743379, 130715]
width = 0.8 # width of the bars
fig, ax = plt.subplots() #make subplots, so that their will be labels above the bars
rects1 = ax.bar(x, y_labels) #rects stands for rectangles, the shape of the bar chart. 

def autolabel(rects):# attach labels above bars #zoals gevonden op matplotlib site
    for rect in rects:
        height = rect.get_height()
        ax.text(rect.get_x() + rect.get_width()/2., 1.05*height,
                '%d' % int(height),
                ha='center', va='bottom') #The placement of the text is determined by the height function, or the height of the column
        #and the number that is put on top of each column is written by: '%d' %int(height). 
        #So all you need to do is create an array of strings that you want at the top of the columns and iterate through. 
#print(autolabel(rects1))

plt.bar(x,y) #plot a bar chart and attach some information to it
plt.title("total number of characters in a volume")
plt.xlabel("volume")
plt.ylabel("# of characters")
plt.xticks(x,x_labels)
#plt.savefig('characters_per_volume.png')
#plt.show()

#visualise the lines per volume
#make a list of the total lines per volume ---> lijst maken van de counters die def(calculate_lines_II(volume)) teruggeeft 
lines_per_volume = []
for volume in corpus: 
	lines_per_volume.append((calculate_lines_II(volume))[2]) 

x = [1,2,3,4,5,6,7,8,9,10]
y = lines_per_volume
x_labels = ["vol1","vol2","vol3","vol4","vol5","vol6","vol7","vol8","vol9","vol10"]
y_labels = [12498, 12159, 13092, 10925, 14585, 10491, 13786, 13189, 12517, 2221 ]
width = 0.8 # width of the bars
fig, ax = plt.subplots() #make subplots, so that their will be labels above the bars
rects2 = ax.bar(x, y_labels) #rects stands for rectangles, the shape of the bar chart. 

def autolabel(rects):# attach labels above bars as recommanded by the matplotlib site
    for rect in rects:
        height = rect.get_height()
        ax.text(rect.get_x() + rect.get_width()/2., 1.05*height,
                '%d' % int(height),
                ha='center', va='bottom') #The placement of the text is determined by the height function, or the height of the column
        #and the number that is put on top of each column is written by: '%d' %int(height). 
        #So all you need to do is create an array of strings that you want at the top of the columns and iterate through. 
#print(autolabel(rects2))

plt.bar(x,y) #plot a bar chart and attach some information to it
plt.title("total number of lines in a volume")
plt.xlabel("volume")
plt.ylabel("# of lines")
plt.xticks(x,x_labels)
#plt.savefig('lines_per_volume.png')
#plt.show()

#visualise the sentences per volume
sentences_per_volume = []
for item in volumes.fileids():
	sentences_per_volume.append(len(volumes.sents(item)))
#print(sentences_per_volume)

x = [1,2,3,4,5,6,7,8,9,10]
y = sentences_per_volume
x_labels = ["vol1","vol2","vol3","vol4","vol5","vol6","vol7","vol8","vol9","vol10"]
y_labels = sentences_per_volume
width = 0.8 # width of the bars
fig, ax = plt.subplots() #make subplots, so that their will be labels above the bars
rects1 = ax.bar(x, y_labels) #rects stands for rectangles, the shape of the bar chart. 

def autolabel(rects):# attach labels above bars as recommanded by the matplotlib site
    for rect in rects:
        height = rect.get_height()
        ax.text(rect.get_x() + rect.get_width()/2., 1.05*height,
                '%d' % int(height),
                ha='center', va='bottom') #The placement of the text is determined by the height function, or the height of the column
        #and the number that is put on top of each column is written by: '%d' %int(height). 
        #So all you need to do is create an array of strings that you want at the top of the columns and iterate through. 
#print(autolabel(rects1))

plt.bar(x,y) #plot a bar chart and attach some information to it
plt.title("total number of sentences in a volume")
plt.xlabel("volume")
plt.ylabel("# of sentences")
plt.xticks(x,x_labels)
#plt.savefig('sentences_per_volume.png')
#plt.show()


#visualise the words per volume
words_per_volume = []
for item in volumes.fileids():
	words_per_volume.append(len(volumes.words(item)))
#print(words_per_volume)

x = [1,2,3,4,5,6,7,8,9,10]
y = words_per_volume
x_labels = ["vol1","vol2","vol3","vol4","vol5","vol6","vol7","vol8","vol9","vol10"]
y_labels = words_per_volume
width = 0.8 # width of the bars
fig, ax = plt.subplots() #make subplots, so that their will be labels above the bars
rects1 = ax.bar(x, y_labels) #rects stands for rectangles, the shape of the bar chart. 

def autolabel(rects):# attach labels above bars as recommanded by the matplotlib site
    for rect in rects:
        height = rect.get_height()
        ax.text(rect.get_x() + rect.get_width()/2., 1.05*height,
                '%d' % int(height),
                ha='center', va='bottom') #The placement of the text is determined by the height function, or the height of the column
        #and the number that is put on top of each column is written by: '%d' %int(height). 
        #So all you need to do is create an array of strings that you want at the top of the columns and iterate through. 
#print(autolabel(rects1))

plt.bar(x,y) #plot a bar chart and attach some information to it
plt.title("total number of words in a volume")
plt.xlabel("volume")
plt.ylabel("# of words")
plt.xticks(x,x_labels)
#plt.savefig('words_per_volume.png')
#plt.show()

#visualise all the total numbers for the entire corpus (= ten volumes of The Arabian Nights)
'''import texttable as tt#import texttable module
tab = tt.Texttable() #initialize texttable object

header = ['Corpus', 'Total characters', 'Total lines', 'Total sentences', 'Total words']#To insert a header we create a list with each element containing the title of a column
tab.header(header)# add it to the table using the header() method of the TextTable object.

row = ['The Arabian Nights', '6627422', '115463', '43094','1459332']
tab.add_row(row) #To insert a row into the table, we create a list with the elements of the row and add it to the table

tab.set_cols_width([18,18,18,18,18])#set the width of the table cells
#set the horizontal and vertical alignment of data within table cells
tab.set_cols_align(['c','c', 'c','c', 'c']) #‘c’ for center alignment 

tab.set_deco(tab.HEADER | tab.VLINES)#control drawing of lines between rows and columns and between the header and the first row
#I choose for a line below the header and lines between columns

tab.set_chars(['-','|','+','#'])#list of elements which determine character used for horizontal lines, vertical lines,
#intersection points of these lines and the header line, in that order

table_statistics = tab.draw()#table is returned as a string
#print(table_statistics)'''

# Now we create a Word document which contains the graphs, so we have an overview
from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading('Visualisation of the statistics')
document.add_paragraph('Overview of the characters per volume')
document.add_picture('characters_per_volume.png', width=Inches(5.25))
document.add_page_break()
document.add_paragraph('Overview of the words per volume')
document.add_picture('words_per_volume.png', width=Inches(5.25))
document.add_paragraph('Overview of the sentences per volume')
document.add_picture('sentences_per_volume.png', width=Inches(5.25))
document.add_paragraph('Overview of the lines per volume')
document.add_picture('lines_per_volume.png', width=Inches(5.25))
document.save('Statistics.docx')


##################################
#Separate the volumes into nights
##################################
'''
# Now that we know some of the statistics about each volume and the entire corpus, we
# want to have a look at the individual nights. 

# To make sure all the volumes are automatically read. 
read_corpus = []
for volume in corpus:
	f = open(volume, 'rt', encoding='utf-8') 
	text = f.read()
	f.close()
	read_corpus.append(text)

import re

# Two definitions that can find the starting and ending index of each night. 
def start_idex_nights(regex, text, flags=re.IGNORECASE): # So it is case insensitive.
	start_index = []
	for match in re.finditer(regex, text, flags):
		start_index.append(match.start())
	return start_index	

def start_and_end_index_nights(volume):
	start_index = start_idex_nights("When it was the.*Night,?\n", volume)
	night_indexes = []
	for i in range(len(start_index)-1): # Minus 1 because the last night of each volume is no new start index of the next night (which is in the next volume).
		night_indexes.append((start_index[i], start_index[i+1])) # You get the index of a night and the subsequent night, which is the end index of the previous night.	
	night_indexes.append((start_index[-1], len(volume) - 1)) # Here the indexes of the last night of a volume are added.
	return night_indexes

counter = 0
for volume in read_corpus:
	counter+= (len(start_and_end_index_nights(volume)))
#print(counter) # It seems there are only 990 nights in the ten volumes, or at least, 990 nights are extracted. 

#Now we put each night into a separate file using the indexes calculated above.

pattern = re.compile(r'\s[Nn]ight')
for volume in read_corpus:
	indexes_night = start_and_end_index_nights(volume)
	for i in indexes_night:
		sentence = volume[i[0]+12:i[0]+150] #150 is just a random number, it makes sure that the first sentence (starting from the number e.g. 'second') is included in 'sentence'
		sentence = pattern.split(sentence)	
		filename = 'data/'+ str(sentence[0]) + '.txt'
		f = open(filename,'wt', encoding='utf-8')
		f.write(volume[i[0]:i[1]])
		f.close()

#I now created 990 files and each file has a name like 'the second', 'the the Four Hundred and Sixty-third',...
#This is important so that we know which night we are talking about.

#####################################
#Calculate statistics for each night
#####################################   	

#Make a new corpus, consisting of the nights so statistics can be calculated
#Each time, I will make a dictionary so it is easy to look up how many words/ lines/ characters... a night has
#Each time, I will also have to make a list. These will then be used to visualise the statistics. 

corpus_root= 'data'
corpus_nightsII = PlaintextCorpusReader(corpus_root, '[Tt]he\s.*')	
print(len(corpus_nightsII.fileids())) #to check whether all 990 nights are in the corpus

#How many characters does each night have? 
import collections
def calculate_characters_nights(corpus):
	characters_per_night = {}
	characters_per_night_list = []
	for night in corpus:
		f = open('data/' + night, 'rt', encoding='utf-8') 
		text = f.read()
		f.close()
		characters_per_night[night] = len(text)
		characters_per_night_list.append((night,len(text)))
		characters_per_night = collections.OrderedDict(characters_per_night) #we make sure that the order of the data stays the same
	return characters_per_night, characters_per_night_list

char_dict_night = (calculate_characters_nights(corpus_nightsII.fileids()))[0]
char_list_night = (calculate_characters_nights(corpus_nightsII.fileids()))[1]

#Which night has the most characters? 
for file, characters in char_list_night:
    if characters == max(char_dict_night.values()):
        print(file, characters) # the Eight Hundred and Forty-fifth.txt => 51841

#How many lines does each night have?
def calculate_lines_night(file):
	count = 0
	f = open ('data/' + file, 'rt', encoding='utf-8') 
	for line in f:
		count += 1
	f.close()
	return count

line_list_nights = []
line_dict_nights = {}
for night in corpus_nightsII.fileids():
	line_list_nights.append((night, calculate_lines_night(night)))
	line_dict_nights[night] = calculate_lines_night(night)
	line_dict_nights = collections.OrderedDict(line_dict_nights) #we make sure that the order of the data stays the same

#Which night has the most lines? 
for file, characters in line_list_nights:
    if characters == max(line_dict_nights.values()):
        print(file, characters) #the Eight Hundred and Forty-fifth.txt => 864

word_dic_nights = {}
word_list_nights = []
for file in corpus_nightsII.fileids(): #calculate the amount of words in each volume
	word_list_nights.append((file, len(corpus_nightsII.words(file))))
	word_dic_nights[file] = len(corpus_nightsII.words(file))
	word_dic_nights = collections.OrderedDict(word_dic_nights) #we make sure that the order of the data stays the same

#Which night has the most words?
for file, characters in word_list_nights:
	if characters == max(word_dic_nights.values()):
		print(file, characters) #the Eight Hundred and Forty-fifth.txt => 11794

sentence_list_nights = []
sentence_dic_nights = {}
for file in corpus_nightsII.fileids(): #calculate the amount of sentences in each volume
	sentence_list_nights.append((file, len(corpus_nightsII.sents(file))))
	sentence_dic_nights[file] = len(corpus_nightsII.sents(file))
	sentence_dic_nights = collections.OrderedDict(sentence_dic_nights) #we make sure that the order of the data stays the same 

#Which night has the most sentences?
for file, characters in sentence_list_nights:
	if characters == max(sentence_dic_nights.values()):
		print(file, characters)	#the Eight Hundred and Forty-fifth.txt => 399

#In the following block of code, we calculate what the average word length is in each night
from nltk import FreqDist
import nltk
dict_word_length = {}
for file in corpus_nightsII.fileids():
	text = corpus_nightsII.words(file)
	x = [len(words) for words in text]
	fdist = FreqDist(x)
	dict_word_length[file] = fdist.max()
#print(dict_word_length)

# We now calculate the readability for each file. We do this by using the Automated Readability Index (ARI).
stat_list = []
x = word_dic_nights.keys()
for name in x:
	n_char = char_dict_night[name]
	n_words = word_dic_nights[name]
	n_sents = sentence_dic_nights[name]
	stat_list.append((name, n_char, n_words, n_sents))
print(stat_list)

def ARI(n_char, n_words, n_sents):
	x = n_char/n_words
	y = n_words/n_sents

	ARI_score = 4.71*x + 0.5*y - 21.43

	return ARI_score

ARI_list = []
ARI_dic = {}
for tuples in stat_list:
	ARI_score = ARI(tuples[1], tuples[2], tuples[3])
	ARI_score = round(ARI_score, 2)
	ARI_dic[tuples[0]] = ARI_score
	ARI_list.append((tuples[0], ARI_score))

# Which night is the most/ least difficult text to read according to the readability-score?
for night, ARI_score in ARI_list:
	if ARI_score == max(ARI_dic.values()):
		print('The night with the highest ARI_score is:', night,':', ARI_score)
	if ARI_score == min(ARI_dic.values()):
		print('The night with the lowest ARI_score is:', night,':', ARI_score)		

#####################################
#visualise statistics for each night
##################################### 

#collect data for table with total numbers for each night

characters_per_night = [] #list of the characters per night
for tuples in char_list_night:
	characters_per_night.append(tuples[1])	
print(characters_per_night)

# Now we have this list, we can also calculate the average amount of characters in a file:
avarage_characters = sum(characters_per_night)/len(characters_per_night)
print(round(avarage_characters)) # On average, there are 6603 characters in a file

lines_per_night = [] #list of the lines per night
for tuples in line_list_nights:
	lines_per_night.append(tuples[1])
print(lines_per_night)

avarage_lines = sum(lines_per_night)/len(lines_per_night)
print(round(avarage_lines)) #On average, there are 166 characters in a file

sentences_per_night = [] #list of the sentences per night
for tuples in sentence_list_nights:
	sentences_per_night.append(tuples[1])
#print(sentences_per_night)

avarage_sentences = sum(sentences_per_night)/len(sentences_per_night)
print(round(avarage_sentences)) #On average, there are 43 sentences in a file

words_per_night = [] #list of the words per night
for tuples in word_list_nights:
	words_per_night.append(tuples[1])
#print(words_per_night)

avarage_words = sum(words_per_night)/len(words_per_night)
print(round(avarage_words)) #On average, there are 1454 words in a file.

number_nights = [] #create a list with the names of the nights, this is in the order from the dictionaries so not chronological
for name in corpus_nightsII.fileids():
	number_nights.append(name[:-4])#remove the extension from the file name
print(number_nights) #Here we discovered that there is a fault in a file. The first line of the story is 'The Hundred and and night', so we actually don't know which number it is

#create table with all the data
import numpy as np
import pandas as pd #import panda so we can turn the data into a data table with pandas dataframe
column1 = number_nights
column2 = characters_per_night
column3 = lines_per_night
column4 = sentences_per_night
column5 = words_per_night

df = pd.DataFrame({'Nights': column1,'Total characters': column2,'Total lines': column3, 'Total sentences': column4,'Total words': column5})
#print(df) #show the data frame

from pandas import ExcelWriter as xlwt #import excelwriter module
from xlwt import Workbook
writer = xlwt('table of all nights.xlsx') #create an excel file from the data frame
workbook = writer.book #define the excel workbook
df.to_excel(writer, 'Sheet1') #place the data frame on the first sheet of the excel file
worksheet = writer.sheets['Sheet1'] #define the worksheet
worksheet.set_column('B:F',35) #set the column width for columns B up to F, so we can see all the text in the cells
format = workbook.add_format({'bold': True, 'font_color': 'red'}) #put in bold and red the night with the highest number of char., lines, sents. and words

#find the highest numbers per column (per list)
print(max(characters_per_night))
print(max(lines_per_night))
print(max(sentences_per_night))
print(max(words_per_night))

worksheet.write('C31', '51841' , format) #highlight the night with the highest number of characters
worksheet.write('D31', '864' , format) #highlight the night with the highest number of lines
worksheet.write('E31', '399' , format) #highlight the night with the highest number of sentences
worksheet.write('F31', '11794' , format) #highlight the night with the highest number of words
worksheet.write('B31', 'the Eight Hundred and Forty-fifth', format)#highlight the name of the night with the highest numbers
writer.save() #save and close the excel file

#now the general visualisations of the statistics are finished, we can start preparing the texts for topic modeling '''

#######################################
#Prepare the texts for topic modeling
#######################################
# First, we make a new corpus consisting of the nights and some additional texts, because we need enough data to apply topic modelling.

'''pattern = re.compile(r'[Tt]he') 
corpus_tales = []
for file in listdir('data'):
	if pattern.search(file):
		corpus_tales.append('data' + '/' + file)
print(len(corpus_tales)) #for topic modeling we should have a minimum of 1000 files, now we have 1030 files (990 nights and some additional tales)

#Now that we have our corpus, we need to tokenize every file so we can leave out the punctuation, stopwords, words that occur only once
# modals, cardinal numbers... and save them in 'clean_doc'.
import string
punc = string.punctuation #import a list of punctuation
from nltk.corpus import stopwords #import a list of stopwords
stoplist = stopwords.words('english')
additional_stopwords = ["thou", "thee", 'thy', "'s"] #these are middle English stopwords that are not in the nltk list and the possesive 's'
additional_punc = ['``','--', "''"] #this is punctuation that might be in some tales, but is not in the punctuation list of nltk
import nltk #import nltk to be able to use the tokenizer
from nltk.stem.porter import PorterStemmer
p_stemmer = PorterStemmer()
from collections import defaultdict
frequency = defaultdict(int)# make an empty default dict so we can compute the frequency of the words and delete words that only occur once

for tale in corpus_tales:
	filtered_text = []
	f = open(tale,'rt', encoding='utf-8')
	text = f.read()
	f.close()
	text = text.lower()
	text = nltk.word_tokenize(text)
	for item in text:
		if item in punc:
			continue
		if item in additional_punc:
			continue	
		if item in stoplist:
			continue
		if item in additional_stopwords:
			continue
		filtered_text.append(item)
	pos_text = nltk.pos_tag(filtered_text)
	for tuples in pos_text: #this will remove modals and cardinal numbers
		if tuples[1] == 'MD':
			filtered_text.remove(tuples[0])
		if tuples[1] == 'CD':
			filtered_text.remove(tuples[0])	
	#filtered_text = [p_stemmer.stem(i) for i in filtered_text] #the words are stemmed. We noticed that some words turned into a rather strange stem and we don't know whether that is how it should be.				
	for words in filtered_text:
		frequency[words] += 1
	filtered_text = [words for words in filtered_text if frequency[words] > 1] #now only words that occur more than once are in 'filtered_text'	
	filename = 'clean_doc/' + str(tale[5:-4]) + '_filtered' + '.txt' # 5:-4 so 'data/' is left out as well as '.txt'.
	f_out = open(filename,'wt', encoding='utf-8')
	f_out.write(' '.join(filtered_text))
	f_out.close()


# Now we make a new corpus consisting of the filtered texts
import re #nog weg doen
pattern = re.compile(r'[Tt]he') 
clean_corpus= []
for file in listdir('clean_doc'):
	if pattern.search(file):
		clean_corpus.append('clean_doc' + '/' + file)
#print(len(clean_corpus)) #To check whether all 1030 files are in the corpus. It is indeed correct.

# Now we make a nested list and afterwards a dictionary, this is necessary for creating the document matrix
import gensim
from gensim import corpora

nested_list = []
for tale in clean_corpus:
	f = open(tale,'rt', encoding='utf-8')
	text = f.read()
	f.close()
	text = nltk.word_tokenize(text)
	nested_list.append(text)
	dictionary = corpora.Dictionary(nested_list)
#dictionary.save('clean_files_dic.txtdic')
#print(dictionary.token2id)


# We are ready to turn the dictionary into a document-term matrix
# Now we convert the dictionary into a bag of words and call it a vector corpus
vector_corpus = [dictionary.doc2bow(text) for text in nested_list] # this gives us the document-term matrix
#print(vector_corpus [2]) #list of sparse vectors equal to the number of documents. 
#In each document the sparse vector is a series of tuples.The tuples are (term ID, term frequency) pairs.

#######################################
# Ready for topic modeling
#######################################

import numpy as np

np.random.seed(1) #setting random seed to get the same results each time.
ldamodel = gensim.models.LdaModel(vector_corpus, num_topics=100, id2word = dictionary, passes=5)
# first parameter: determine how many topics should be generated. Our document set is relatively large, so we’re  asking for 100 topics.
# second parameter: our previous dictionary to map ids to strings
# third parameter: number of laps the model will take through corpus. More passes = more accurate model. 
#But a lot of passes can be slow on a very large corpus.So let's say we do 5 laps.

#ldamodel.save('topicmodel.lda') #We save and load the model for later use instead of having to rebuild it every time
ldamodel = gensim.models.LdaModel.load('topicmodel.lda')

#print(ldamodel.show_topics(num_topics=-1, num_words=4)) #prints the num_words most probable words for all topics to log. topics=-1 to print all topics.
# first parameter defines the number of topics, second parameter the number of words per topic, this is 10 words per topic by default
#print(ldamodel.print_topics(5) #print the most contributing words for ... randomly selected topics

# Now that we have our ldamodel and have an idea about the topics that are in fairy tales, we want to test the model
# on our original corpus of tales: corpus_nightsII (in order to do that, we need to convert in into a BOW representation)

# We are not sure whether we need to clean the testcorpus as well. We did it to be sure.

for night in corpus_nightsII.fileids():
	filtered_text = []
	f = open('data/' + night, 'rt', encoding='utf-8') 
	text = f.read()
	f.close()
	text = text.lower()
	text = nltk.word_tokenize(text)
	for item in text:
		if item in punc:
			continue
		if item in additional_punc:
			continue	
		if item in stoplist:
			continue
		if item in additional_stopwords:
			continue
		filtered_text.append(item)
	pos_text = nltk.pos_tag(filtered_text)
	for tuples in pos_text: #this will remove modals and cardinal numbers
		if tuples[1] == 'MD':
			filtered_text.remove(tuples[0])
		if tuples[1] == 'CD':
			filtered_text.remove(tuples[0])	
	#filtered_text = [p_stemmer.stem(i) for i in filtered_text] #the words are stemmed. We noticed that some words turned into a rather strange stem and we don't know whether that is how it should be.				
	for words in filtered_text:
		frequency[words] += 1
	filtered_text = [words for words in filtered_text if frequency[words] > 1] #now only words that occur more than once are in 'filtered_text'	
	filename = 'clean_nights/' + str(night[:-4]) + '_filtered' + '.txt'
	f_out = open(filename,'wt', encoding='utf-8')
	f_out.write(' '.join(filtered_text))
	f_out.close()

# The cleaned nights are now in 'clean_nights', now we can convert them to a BOW representation	
pattern = re.compile(r'[Tt]he') 
clean_nights_corpus = []
for file in listdir('clean_nights'):
	if pattern.search(file):
		clean_nights_corpus.append('clean_nights' + '/' + file)

corpus = []
for file in clean_nights_corpus:
	f = open(file,'rt', encoding='utf-8')
	text = f.read()
	f.close()
	text = nltk.word_tokenize(text)
	bow_vector = dictionary.doc2bow(text)
	#lda_vector = ldamodel[bow_vector]
	lda_vector = ldamodel.get_document_topics(bow_vector, minimum_probability=0.0)
	lda_vector = [b for (a, b) in sorted(lda_vector)]
	corpus.append(lda_vector)

# Now we would like to print every document's single most prominent LDA topic in a separate txt file
f_out = open('topic_per_night.txt','at', encoding='utf-8')
for file in clean_nights_corpus:
	f = open(file,'rt', encoding='utf-8')
	text = f.read()
	f.close()
	text = nltk.word_tokenize(text)
	bow_vector = dictionary.doc2bow(text)
	lda_vector = ldamodel[bow_vector]
	topic = ldamodel.print_topic(max(lda_vector, key=lambda item: item[1])[0])
	f_out.write('\n' + file + ':\n' + topic + '\n')
f_out.close()	

# Now we make a matrix of the documents & topics:
X = ldamodel.show_topics(num_topics= 100, num_words=10) #not necessary, was just a test to see if it made any difference
X = np.array(corpus) #should be the matrix containing the nights & the topics'''

#print(X.shape) to check if the shape of our matrix is suitable for hierarchical clustering, it shoudl give the number of files and number of topics

#########################
# Evaluate our LDA model
#########################
# We will split each document into two parts, and check that topics of the first half are similar to topics 
# of the second half. And we check whether two halves of different documents are less similar.
# The halves of the same document should be very similar, the halves of different documents should be a bit 
# less similar, although similarity is expected here as well, since they are all fairy tales.
'''def intra_inter(model, test_docs, num_pairs=2000):
    part1 = []
    part2 = []
    for file in test_docs:
    	f = open(file, 'rt', encoding='utf-8') 
    	text = f.read()
    	f.close()
    	text = nltk.word_tokenize(text)# split each test document into two halves and compute topics for each half
    	x = round(len(text)/ 2)
        part1.append(model[dictionary.doc2bow(text[:x])])
        part2.append(model[dictionary.doc2bow(text[x:])])
    
    # print computed similarities (uses cossim)
    print('average cosine similarity between corresponding parts (higher is better):')
    print(np.mean([gensim.matutils.cossim(p1, p2) for p1, p2 in zip(part1, part2)]))

    random_pairs = np.random.randint(0, len(test_docs), size=(num_pairs, 2))
    print('average cosine similarity between 2000 random parts (should be a bit lower):')    
    print(np.mean([gensim.matutils.cossim(part1[i[0]], part2[i[1]]) for i in random_pairs]))

print(intra_inter(ldamodel, clean_nights_corpus))   

################
#Topic richness
################
from nltk import FreqDist #verwijderen als alles 'aan' staat
# This code will compute how many topics are assigned to each file
number_of_topics = {} # we make a dictionary so that it is easy to look up the file and its number of topics
number_of_topics_list = [] #we also make a list which comes in handy to calculate which file has the highest amount of topics
for file in clean_nights_corpus:
	f = open(file,'rt', encoding='utf-8')
	text = f.read()
	f.close()
	text = nltk.word_tokenize(text)
	bow_vector = dictionary.doc2bow(text)
	lda_vector = ldamodel[bow_vector]
	name = file[13:-13] #filter out 'clean_nights/' and '_filtered.txt'
	number_of_topics[name] = len(lda_vector)
	number_of_topics_list.append((name, len(lda_vector)))

for file, topics in number_of_topics_list:
	if topics == max(number_of_topics.values()):
		print(file, ' has the highest amount of topics: ', topics)
	if topics == min(number_of_topics.values()):
		print(file, ' has the lowest amount of topics: ', topics)

# We also want to now what the average number of topics is
from statistics import mean
print('The average number of topics is', int(mean(number_of_topics.values())))

# In order to know what number of topics occurs most often, we make a FreqDist:
fdist_topics = FreqDist(number_of_topics.values())	
#print(fdist_topics.max(), ' is the number of topics that occurs most often')	

# We also want a top 10 of the files, the files with the highest number of topics

from collections import Counter
top_10 = (dict(Counter(number_of_topics).most_common(10)))	
#print(top_10) 

# Create a table of the file with the maximum and minimum number of topics
import pandas as pd #nog weg doen op het einde

# First we need the lda vector of the night with the maximum number of topics
f = open('clean_nights/the Three Hundred and Fifty-seventh_filtered.txt','rt', encoding='utf-8') 
text = f.read()
f.close()
text = nltk.word_tokenize(text)
bow_vector = dictionary.doc2bow(text)
lda_vector_mintop = ldamodel[bow_vector]
#print(lda_vector_mintop)

# Then we need the lda vector of the night with the minimum number of topics
f = open('clean_nights/the Six Hundred and Sixty-second_filtered.txt','rt', encoding='utf-8')
text = f.read()
f.close()
text = nltk.word_tokenize(text)
bow_vector = dictionary.doc2bow(text)
lda_vector_maxtop = ldamodel[bow_vector]
#print(lda_vector_maxtop)

def final_topics(lda_vector):
	x = []
	for topics in lda_vector:  # get the topics for the maximum topics file and the minimum topics file
		x.append(ldamodel.show_topic(topics[0]))
		
	y = []
	z = []
	for list1 in x:
		for tuples in list1:
			y.append(tuples[0]) # add the words belonging tho the same topic to a list
		z.append(y) # add that list to z
		y = []	# empty list y again for the next topic
		

		final_topic_list = []
		for topic in z:
			topic = topic [:5] # now we get only the first five words per topic
			topic = ' '.join(topic)
			final_topic_list.append(topic)
	return final_topic_list
max_top = (final_topics(lda_vector_maxtop))
min_top = (final_topics(lda_vector_mintop))
print(len(max_top)) # to check if it matches the number of topics in the table we will make later on
print(len(min_top))

maxmin_topics = [] #make a nested list with as first list max_top and second list min_top
maxmin_topics.append(max_top)
maxmin_topics.append(min_top) 


#prepare the data for the data frame
column1 = ['max number of topics', 'min number of topics' ]
column2 = ['The Six Hundred and Sixty-second', 'The Three Hundred and Fifty-seventh']
column3 = ['16', '5']
column4 = maxmin_topics

df = pd.DataFrame({'Max/Min topics': column1,'Nights': column2,'Number of topics': column3, 'Topics': column4})
print(df) #show the data frame

from pandas import ExcelWriter as xlwt #nog weg doen
from xlwt import Workbook #nog weg doen
writer = xlwt('table of max and min number topics.xlsx') #create an excel file from the data frame
workbook = writer.book #define the excel workbook
df.to_excel(writer, 'Sheet1') #place the data frame on the first sheet of the excel file
worksheet = writer.sheets['Sheet1'] #define the worksheet
worksheet.set_column('B:Q',35) #set the column width for columns B up to Q, so we can see all the text in the cells'''
#writer.save()

###########################################
# Hierarchical clustering with topic model
###########################################
'''from matplotlib import pyplot as plt   #this should be some start code for the hierarchical clustering of our topics. Not finished yet!
import numpy as np

from scipy.spatial.distance import pdist, squareform
dm = squareform(pdist(X, 'cosine'))#'cosine'is one of the methods that can be used to calculate the distance between newly formed clusters
								   #we use the cosine similarity because it works good for topic clustering

from scipy.cluster.hierarchy import linkage #creating a linkage matrix
linkage_object = linkage(dm, method='ward', metric='euclidean')
#print(linkage_object) #linkage_object[i] will tell us which clusters were merged in the i-th pass


#calculate a full dendrogram
from scipy.cluster.hierarchy import dendrogram 
plt.figure(figsize=(25, 10))
plt.title('Hierarchical Clustering Dendrogram')
plt.xlabel('sample index')
plt.ylabel('distance')
dendrogram(linkage_object,leaf_rotation=90.,leaf_font_size=8.,)
plt.show()

#we create a truncated dendrogram, which only shows the last p=15 out of our 989 merges.
from scipy.cluster.hierarchy import dendrogram
plt.title('Hierarchical Clustering Dendrogram (truncated)') 
plt.xlabel('sample index')
plt.ylabel('distance')
dendrogram(linkage_object,truncate_mode='lastp',p=15,show_leaf_counts=False,leaf_rotation=90.,leaf_font_size=12.,show_contracted=True)
#lastp means 'show only the last p merged clusters',  show_leaf_counts=False, because otherwise numbers in brackets are counts,
#show_contracted=True, because we want to get a distribution impression in truncated branches
plt.show()

#now we create a 'fancy dendrogram' by annotating the distances inside the dendrogram 
def fancy_dendrogram(*args, **kwargs):
    max_d = kwargs.pop('max_d', None)
    if max_d and 'color_threshold' not in kwargs:
        kwargs['color_threshold'] = max_d
    annotate_above = kwargs.pop('annotate_above', 0)

    ddata = dendrogram(*args, **kwargs)

    if not kwargs.get('no_plot', False):
        plt.title('Hierarchical Clustering Dendrogram (truncated and fancy)')
        plt.xlabel('sample index or (cluster size)')
        plt.ylabel('distance')
        for i, d, c in zip(ddata['icoord'], ddata['dcoord'], ddata['color_list']):
            x = 0.5 * sum(i[1:3])
            y = d[1]
            if y > annotate_above:
                plt.plot(x, y, 'o', c=c)
                plt.annotate("%.3g" % y, (x, y), xytext=(0, -5),
                             textcoords='offset points',
                             va='top', ha='center')
        if max_d:
            plt.axhline(y=max_d, c='k')
    return ddata

fancy_dendrogram(Z,truncate_mode='lastp', p=15,leaf_rotation=90.,leaf_font_size=12., show_contracted=True,annotate_above=10.)
plt.show()

#Get the number of flat clusters from a linkage matrix at a specified distance.
def num_clusters(hc, d):
	return len(numpy.unique(scipy.cluster.hierarchy.fcluster(linkage_object, 30, criterion='distance')))#d (number): Distance threshold for defining flat clusters.

number_clusters = num_clusters(linkage_object, 30)
print(number_clusters) #this does not work, i think this is not the right way to print it

###############
#Visualisation
###############
import matplotlib.pyplot as plt
import matplotlib.patches as patches

# This code will iterate over the dictionary and print a random set of words in the color of the topic 
# it belongs to most. 

# These are colors assigned to the hundred topics.
topic_colors = {0:'#FFC400', 1:'#30a2da', 2:'#FFFAF0', 3:'#FFC0CB', 4:'#B0E0E6', 5:'#FF0000', 6:'#FAA460', 7:'#2E8B57', 
                8:'#FFF5EE', 9:'#A0522D', 10:'#C0C0C0', 11:'#87CEEB', 12:'#00FFFF', 13:'#7FFFD4', 
                14:'#F0FFFF', 15:'#F5F5DC', 16:'#FFE4C4', 17:'#000000', 18:'#FFEBCD', 19:'#0000FF', 20:'#8A2BE2',
                21:'#A52A2A', 22:'#DEB887', 23:'#5F9EA0', 24:'#7FFF00', 25:'#D2691E', 26:'#FF7F50', 27:'#6495ED', 
                28:'#DC143C', 29:'#00FFFF', 30:'#00008B', 31:'#008B8B', 32:'#B8860B', 33:'#A9A9A9', 34:'#006400',
                35:'#BDB76B', 36:'#8B008B', 37:'#556B2F', 38:'#FF8C00', 39:'#9932CC', 40:'#E9967A', 41:'#8FBC8F',
                42:'#483D8B', 43:'#2F4F4F', 44:'#00CED1', 45:'#9400D3', 46:'#FF1493', 47:'#00BFFF', 48:'#696969',
                49:'#1E90FF', 50:'#B22222', 51:'#228B22', 52:'#FF00FF', 53:'#DCDCDC', 54:'#F8F8FF', 55:'#FFD700',
                56:'#DAA520', 57:'#808080', 58:'#008000', 59:'#ADFF2F', 60:'#F0FFF0', 61:'#FF69B4', 62:'#CD5C5C',
                63:'#4B0082', 64:'#FFFFF0', 65:'#F0E68C', 66:'#E6E6FA', 67:'#FFF0F5', 68:'#7CFC00', 69:'#FFFACD',
                70:'#ADD8E6', 71:'#F08080', 72:'#E0FFFF', 73:'#FAFAD2', 74:'#90EE90', 75:'#D3D3D3', 76:'#FFB6C1',
                77:'#FFA07A', 78:'#20B2AA', 79:'#87CEFA', 80:'#778899', 81:'#B0C4DE', 82:'#FFFFE0', 83:'#00FF00',
                84:'#32CD32', 85:'#FAF0E6', 86:'#FF00FF', 87:'#800000', 88:'#66CDAA', 89:'#0000CD', 90:'#BA55D3',
                91:'#9370DB', 92:'#3CB371', 93:'#7B68EE', 94:'#00FA9A', 95:'#48D1CC', 96:'#C71585', 97:'#191970',
                98:'#F5FFFA', 99:'#FFE4E1'}

def color_words_dictionary(model, dictionary):
    fig = plt.figure()
    ax = fig.add_axes([0,0,1,1])

    word_pos = 1/len(dictionary)

    for word in dictionary:
        list1 = model.get_term_topics(word)
        test_list = []
        for tuples in list1:
            test_list.append(tuples[1])
            if tuples[1] == max(test_list):
                x = tuples[0] #choose number of topic that is highest! 
        ax.text(0.2, word_pos, model.id2word[word],
                    horizontalalignment='center',
                    verticalalignment='center',
                    fontsize=20, color=topic_colors[x],  # choose just the most likely topic
                    transform=ax.transAxes)
        word_pos += 0.05

    ax.set_axis_off()
    plt.show()   

#print(color_words_dictionary(ldamodel, dictionary))  

# In this code will iterate over a particular night and print a set of words that are in the file in 
# the color of the topic it belongs to (it also looks at the surrounding words). 
def color_words_night(model, file):
    f = open(file, 'rt', encoding='utf-8') 
    text = f.read()
    f.close()
    text = nltk.word_tokenize(text)
    text = model.id2word.doc2bow(text)
    
    # get word_topics
    doc_topics, word_topics, phi_values = model.get_document_topics(text, per_word_topics=True)

    fig = plt.figure()
    ax = fig.add_axes([0,0,1,1])

    word_pos = 1/len(text)

    for word, topics in word_topics:
        ax.text(0.2, word_pos, model.id2word[word],
                horizontalalignment='center',
                verticalalignment='center',
                fontsize=20, color=topic_colors[topics[0]], 
                transform=ax.transAxes)
        word_pos += 0.2

    ax.set_axis_off()
    plt.show()    

# Say for example, we want to have a look at the Eight Hundred and Eighth night
night = 'clean_nights/the Eight Hundred and Eighth_filtered.txt'
#print(color_words_night(ldamodel, night))

#############################
#next we try to create word clouds, but we had some trouble installing the WordCloud Package, so we weren't able to test this code

#creating square word clouds for our clean_nights corpus
from os import path
from wordcloud import WordCloud
from matplotlib import plt
from PIL import Image
import numpy as np

for file in clean_nights_corpus: #Read the clean_nights corpus
	f = open(file,'rt', encoding='utf-8')
	text = f.read()
	f.close()
	wordcloud = WordCloud().generate(text) # Generate a word cloud image

plt.imshow(wordcloud)# Display the generated image
plt.axis("off")

wordcloud = WordCloud(max_font_size=40).generate(text)
plt.figure()
plt.imshow(wordcloud)
plt.axis("off")
plt.show()

# we make a masked word cloud by using a mask to generate word clouds in arbitrary shapes.
for file in clean_nights_corpus:# Read the clean_nights corpus
	f = open(file,'rt', encoding='utf-8')
	text = f.read()
	f.close()
	skyline_mask = np.array(Image.open("city+skyline.png")))# read the mask image
	wc = WordCloud(background_color="white", max_words=2000, mask=skyline_mask)
	wc.generate(text)# generate word cloud

wc.to_file("city+skyline.png")# store to file

plt.imshow(wc) #make the masked cloud visible
plt.axis("off")
plt.figure()
plt.imshow(skyline_mask, cmap=plt.cm.gray)
plt.axis("off")
plt.show()'''











